from docx import Document
from docx.shared import Pt
import csv


class FamilyDinner:
    def createFamilyDinner(menu, inputFile, outputFile, directoryName):
        document = Document(inputFile)
        with open('menu.csv', mode="r", newline='') as csvfile:
            file = csv.reader(csvfile, delimiter=',')
            for file_row in file:
                # print(file_row)
                name = file_row[0]
                # print("Name:", name)
                price = file_row[1]
                id = file_row[4]
                # print(id)
                
                for table in document.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                if id in paragraph.text:
                                    word = paragraph.text
                                    name = name.strip()[:4]
                                    print("Found", id)

                                    inline = paragraph.runs
                                    # Loop added to work with runs (strings with same style)
                                    for i in range(len(inline)):
                                        print("inline", inline[i].text)
                                        if id in inline[i].text:
                                            text = inline[i].text.replace(id, price)
                                            inline[i].text = text
                                            print("text:", text)

        document.save(directoryName + "/" + outputFile)