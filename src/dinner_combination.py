from docx import Document
from docx.shared import Pt
import csv
import re

class DinnerCombination:
    def createDinnerCombination(menu, inputFile, outputFile, directoryName):
        document = Document(inputFile)
        with open(menu, mode="r", newline='') as csvfile:
            file = csv.reader(csvfile, delimiter=',')

            extraParsingIDs = ["#074", "#075", "#076", "#077"]

            for file_row in file:
                # print(file_row)
                name = file_row[0]
                # print("Name:", name)
                price = file_row[1]
                id = file_row[4]
                # print(id)

                for extraParsingID in extraParsingIDs:
                    if extraParsingID == id:
                        priceList = price.split("&nbsp; &nbsp;")
                        price = ""
                        for i in priceList:
                            i = re.sub("\s+", "", i)
                            print("before i", i)
                            price += i + " "

                        # Adding space before $ sign
                        price = price[:5] + ": " + price[5:16] + ": " + price[16:]

                for table in document.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                if id in paragraph.text:
                                    word = paragraph.text
                                    name = name.strip()[:4]
                                    print("Found", id)


                                    inline = paragraph.runs
                                    # Loop added to work with runs (strings with same style)
                                    for i in range(len(inline)):
                                        print("inline", inline[i].text)
                                        if id in inline[i].text:
                                            text = inline[i].text.replace(id, price)
                                            inline[i].text = text
                                            print("text:", text)

        document.save(directoryName + "/" + outputFile)