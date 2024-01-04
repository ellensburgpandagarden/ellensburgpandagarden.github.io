from docx import Document
from docx.shared import Pt
import csv

class LunchCombination:
    def createLunchCombination(menu, inputFile, outputFile, directoryName):
        document = Document(inputFile)
        with open(menu, mode="r", newline='') as csvfile:
            file = csv.reader(csvfile, delimiter=',')
            for file_row in file:
                # print(file_row)
                name = file_row[0]
                # print("Name:", name)
                price = file_row[1]
                id = file_row[4]
                # print(id)
                
                for table in document.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                if id in paragraph.text:
                                    word = paragraph.text
                                    name = name.strip()[:4]
                                    print("Found", id)

                                    # p = document.add_paragraph(price)

                                    # # Replace element with word
                                    # paragraph.text = paragraph.text.replace(id, price)

                                    inline = paragraph.runs
                                    # Loop added to work with runs (strings with same style)
                                    for i in range(len(inline)):
                                        if id in inline[i].text:
                                            text = inline[i].text.replace(id, price)
                                            inline[i].text = text


                                    # # Format/Style cell
                                    # run = cell.paragraphs[0].runs[0]
                                    # # Fonts used: "Agency FB", "Sitka Banner"
                                    # run.font.name = "Sitka Banner"
                                    # # Fonts size used: "15", "13.5"
                                    # run.font.size = Pt(15) # Change back to 15

        document.save(directoryName + "/" + outputFile)