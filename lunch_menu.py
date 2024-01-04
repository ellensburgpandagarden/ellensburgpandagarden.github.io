from docx import Document
from docx.shared import Pt
import csv
import re

class LunchMenu:
    def createLunchMenu(menu, inputFile, outputFile, directoryName):
        document = Document(inputFile)
        with open(menu, mode="r", newline='') as csvfile:
            file = csv.reader(csvfile, delimiter=',')

            extraParseIDs = ["#204", "#205", "#206", "#207"]

            for file_row in file:
                # print(file_row)
                name = file_row[0]
                # print("Name:", name)
                price = file_row[1]
                id = file_row[4]
            
                # Egg Foo Young extra parsing
                for extraParseID in extraParseIDs:
                    if extraParseID == id:
                        priceList = price.split("&nbsp; &nbsp;")
                        print("Result of Parse: ", priceList)
                        price = ""
                        for i in priceList:
                            i = re.sub("\t+", "", i)
                            print("before i", i)
                            price += i + "          "
                
                extraParsing2 = ["#179", "#180", "#181"]
                # Side Orders (Steamed Rice, Fried Rice, Pan Fried Noodles) extra parsing
                for extraP in extraParsing2:
                    if extraP == id:
                        priceList = price.split("&nbsp; &nbsp;")
                        price = ""
                        for i in priceList:
                            i = re.sub("\s+", "", i)
                            price += i + "  "
                        
                        # Adding space before $ sign
                        price = price[:5] + ": " + price[5:17] + ": " + price[17:]

                for table in document.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                if id in paragraph.text:
                                    word = paragraph.text
                                    name = name.strip()[:4]
                                    # print("Found", id)
                                    # print("Text:", word)

                                    # # p = document.add_paragraph(price)

                                    # # Replace element with word
                                    # paragraph.text = paragraph.text.replace(id, price)

                                    # # Format/Style cell
                                    # run = cell.paragraphs[0].runs[0]
                                    # # Fonts used: "Agency FB", "Sitka Banner"
                                    # run.font.name = "Agency FB"
                                    # # Fonts size used: "15", "13.5"
                                    # run.font.size = Pt(15) # Change back to 15


                                    inline = paragraph.runs
                                    # Loop added to work with runs (strings with same style)
                                    for i in range(len(inline)):
                                        print("inline", inline[i].text)
                                        if id in inline[i].text:
                                            text = inline[i].text.replace(id, price)
                                            inline[i].text = text
                                            print("text:", text)

                                

        document.save(directoryName + "/" + outputFile)